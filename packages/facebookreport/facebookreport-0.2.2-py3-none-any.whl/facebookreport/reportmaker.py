import pandas as pd
import datetime as dt
import matplotlib.pyplot as plt
from statistics import mean
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from urllib.parse import urlencode
from urllib.request import urlretrieve
import sys
from tkinter import *
from tkinter import ttk
from tkinter import filedialog as fd
from tkinter import messagebox as mb
import pickle
import os

class programo(Tk):
    def __init__(self, parent):
        self.parent = parent
        self.iniciar()

    def iniciar(self):

        def chequearArchivo(tipo):
            if tipo == 'post':
                filepath = filepost.get()
            elif tipo == 'page':
                filepath = filepage.get()
            fb = pd.read_csv(filepath, sep=",", skiprows=[1], encoding="latin-1")
            if (tipo == 'post' and 'Lifetime Post Total Reach' in fb.columns ) or (tipo == 'page' and 'Weekly Page Engaged Users' in fb.columns):
                pass
            else:
                error = mb.showerror('Error', 'El archivo ingresado no corresponde al archivo requerido.')
                if tipo == 'post':
                    filepost.set('')
                elif tipo == 'page':
                    filepage.set('')

        def chequearGenerar():
            filepostpath = filepost.get()
            filepagepath = filepage.get()
            listaopcion = lista.get()
            if len(filepostpath) != 0 and len(filepagepath) != 0 and os.path.exists(filepostpath) and os.path.exists(filepagepath) and lista.current() != -1:
                generar['state'] = '!disable'
            else:
                generar['state'] = 'disable'

        def generarInforme():

            futuro = fd.asksaveasfilename(defaultextension = ".docx", filetypes = [("docx", ".docx")], confirmoverwrite=True)

            cuentaName = cuenta.get()

            #Ingreso y edición del archivo csv con estadísticas de la Facebook fanpage
            filepath = filepost.get()
            fb = pd.read_csv(filepath, sep=",", 
                skiprows=[1], encoding="latin-1", 
                converters={
                6: lambda x: dt.datetime.strptime(x,'%m/%d/%Y %I:%M:%S %p'),
                })
            fb.columns = fb.columns.str.replace('\s', '_')
            totalposts = len(fb.index)

            filepath2 = filepage.get()
            fb2 = pd.read_csv(filepath2, sep=",",
                skiprows=[1], encoding="latin-1",
                converters={
                0: lambda x: dt.datetime.strptime(x,'%Y-%m-%d'),
                })
            fb2.columns = fb2.columns.str.replace('\s', '_')

            #Convierte todos los valores inexistentes en cero
            fb = fb.fillna(0)
            fb2 = fb2.fillna(0)

            fb['engagement_rate'] = fb['Lifetime_Engaged_Users'] / fb['Lifetime_Post_Total_Reach']
            fb2['engagement_rate'] = fb2['Daily_Page_Engaged_Users'] / fb2['Daily_Total_Reach']

            base = criterio.get()

            #Sacar el promedio de engagement rate
            mediaER = mean(fb['engagement_rate']) #A nivel post
            mediaERpag = mean(fb2['engagement_rate']) #A nivel page

            #Extracción de información del post más popular
            pd.options.display.max_colwidth = 100
            pop = fb.nlargest(1, base)
            maslink = pop['Permalink']
            mastime = pop['Posted']
            mastype = pop['Type']
            masreach = int(pop['Lifetime_Post_Total_Reach'])
            masimp = int(pop['Lifetime_Post_Total_Impressions'])
            maseng = int(pop['Lifetime_Engaged_Users'])
            masshared = int(pop['Lifetime_Post_Stories_by_action_type_-_share'])
            masER = float(pop['engagement_rate'])

            #Extracción de informacion a nivel página
            likesAhora = fb2['Lifetime_Total_Likes'].iloc[-1]
            newlikes = fb2['Lifetime_Total_Likes'].iloc[-1] - fb2['Lifetime_Total_Likes'].iloc[0]
            usuarios = mean(fb2['Daily_Page_Engaged_Users'])
            alcance = mean(fb2['Daily_Total_Reach'])
            impresiones = mean(fb2['Daily_Total_Impressions'])

            #Definir periodo analizado en el informe
            primerDia = fb2['Date'].iloc[0].to_pydatetime()
            ultimoDia = fb2['Date'].iloc[-1].to_pydatetime()
            if int(primerDia.year) == int(ultimoDia.year):
                if int(primerDia.month) == int(ultimoDia.month):
                    #El tiempo está asignado a partir del tiempo en que
                    # se posteó el post más popular.
                    mesNumero = int(mastime.dt.month)
                    if mesNumero == 1:
                        mes = 'enero'
                    elif mesNumero == 2:
                        mes = 'febrero'
                    elif mesNumero == 3:
                        mes = 'marzo'
                    elif mesNumero == 4:
                        mes = 'abril'
                    elif mesNumero == 5:
                        mes = 'mayo'
                    elif mesNumero == 6:
                        mes = 'junio'
                    elif mesNumero == 7:
                        mes = 'julio'
                    elif mesNumero == 8:
                        mes = 'agosto'
                    elif mesNumero == 9:
                        mes = 'septiembre'
                    elif mesNumero == 10:
                        mes = 'octubre'
                    elif mesNumero == 11:
                        mes = 'noviembre'
                    elif mesNumero == 12:
                        mes = 'diciembre'
                    else:
                        mes = input('Ingresar periodo analizado en el informe\n')
                else:
                    mes = str(int(primerDia.day)) + "/" + str(int(primerDia.month)) +\
                    " al " + str(int(ultimoDia.day)) + "/" + str(int(ultimoDia.month)) +\
                    " del "
                anho = int(mastime.dt.year)
            else:
                mes = str(int(primerDia.day)) + "/" + str(int(primerDia.month)) + "/" +\
                str(int(primerDia.year))+ " al " + str(int(ultimoDia.day)) + "/" +\
                str(int(ultimoDia.month)) + "/" + str(int(ultimoDia.year))
                anho = ""

            #Nombres de los archivos guardados
            screenshot = 'screenshotfacebookpost.png'
            pnghora = 'mejorhora.png'
            pnglikes = 'totallikes.png'
            pngusers = 'engagedusers.png'
            pngreach = 'totalreach.png'
            pngimpres = 'totalimpressions.png'
            pngengage = 'totalengagement.png'

            #Top Ten más populares / Análisis de hora.
            pop2 = fb.nlargest(10, base)
            time = pop2['Posted']
            fighora = time.groupby(pop2['Posted'].dt.hour).count().plot(kind='bar')
            plt.title('Horario de los mejores posts')
            plt.savefig(pnghora)
            plt.close()

            #Otros gráficos
            fig1 = fb2.plot(x='Date', y='Lifetime_Total_Likes')
            plt.title('Total de "me gusta" en el tiempo')
            plt.savefig(pnglikes)
            plt.close()
            fig2 = fb2.plot(x='Date', y='Daily_Page_Engaged_Users')
            plt.title('Usuarios que interactuaron en el tiempo')
            plt.savefig(pngusers)
            plt.close()
            fig3 = fb2.plot(x='Date', y='Daily_Total_Reach')
            plt.title('Alcance total en el tiempo')
            plt.savefig(pngreach)
            plt.close()
            fig4 = fb2.plot(x='Date', y='Daily_Total_Impressions')
            plt.title('Impresiones totales en el tiempo')
            plt.savefig(pngimpres)
            plt.close()
            fig5 = fb2.plot(x='Date', y='engagement_rate')
            plt.title('Tasa de participación en el tiempo')
            plt.savefig(pngengage)
            plt.close()

            #Escritura del documento
            document = Document()
            p = document.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run('Informe {0} {1} {2}'.format(cuentaName, mes, anho)).bold = True

            #Cuenta de Facebook
            p2 = document.add_paragraph()
            p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.add_run('Facebook').bold = True

            p3 = document.add_paragraph()
            p3.add_run('{0} nuevos likes (llegando a {1}).'.format(newlikes, likesAhora))

            p4 = document.add_paragraph()
            p4.add_run('{0} posts emitidos.'.format(totalposts))

            p5 = document.add_paragraph()
            p5.add_run('{0:.2f} usuarios en promedio que realizaron algún tipo de interacción.'.format(usuarios))

            p6 = document.add_paragraph()
            p6.add_run('{0:.2f} personas alcanzadas (promedio de usuarios únicos en cada post).'.format(alcance))

            p7 = document.add_paragraph()
            p7.add_run('{0:.2f} impresiones en promedio.'.format(impresiones))

            p75 = document.add_paragraph()
            p75.add_run('{0:.2f}% tasa de participación (engagement rate) promediada.'.format(mediaER*100))

            p8 = document.add_paragraph('Post destacado:\n')
            #Porque maslink es un pandas Series
            linkpost = str(maslink.values[0])
            p8.add_run('{0}'.format(linkpost))

            #Falta mejorar estas lineas en caso de error: Exceptions management
            params = urlencode({'url':'{0}'.format(linkpost), 'access_key':'fa156802063d446e85884fa21a4a542a'})
            urlretrieve('https://apileap.com/api/screenshot/v1/urltoimage?' + params, screenshot)

            pimage = document.add_picture(screenshot, width=Inches(5))
            #--------------------------------------------
            #INSERTAR ACÁ SCREENSHOT DEL POST MÁS POPULAR
            #--------------------------------------------

            p85 = document.add_paragraph()
            p85.add_run('{0:.2f}% tasa de participación.'.format(masER*100))

            p9 = document.add_paragraph()
            p9.add_run('{0} usuarios únicos alcanzados.'.format(masreach))

            p10 = document.add_paragraph()
            p10.add_run('{0} impresiones.'.format(masimp))

            p11 = document.add_paragraph()
            p11.add_run('{0} interacciones.'.format(maseng))

            if int(masshared) > 0:
                p12 = document.add_paragraph()
                p12.add_run('{0} veces compartido.'.format(masshared))

            document.add_page_break()

            #Sección de gráficos

            pimage2 = document.add_picture(pnglikes, width=Inches(5))
            pimage3 = document.add_picture(pngimpres, width=Inches(5))

            pimage4 = document.add_picture(pngusers, width=Inches(5))
            pimage5 = document.add_picture(pngreach, width=Inches(5))

            pimage6 = document.add_picture(pngengage, width=Inches(5))
            pimage7 = document.add_picture(pnghora, width=Inches(5))

            document.add_page_break()

            #Guardado del documento
            document.save('{0}'.format(futuro))

            #Eliminación de imágenes
            os.remove(pnglikes)
            os.remove(pnghora)
            os.remove(pngengage)
            os.remove(pngreach)
            os.remove(pngusers)
            os.remove(pngimpres)
            os.remove(screenshot)

            #Apertura de archivo
            if sys.platform == 'linux':
                os.system("xdg-open " + futuro)
            else:
                os.startfile(futuro)

        def seleccionArchivo(nivel):
            if nivel == 'post':
                filename = fd.askopenfilename(defaultextension = '.csv', filetypes = [('csv', '*.csv')])
                filepost.set(filename)
                chequearArchivo('post')
                chequearGenerar()
            if nivel == 'page':
                filename = fd.askopenfilename(defaultextension = '.csv', filetypes = [('csv', '*.csv')])
                filepage.set(filename)
                chequearArchivo('page')
                chequearGenerar()

        def confirmarEliminar(nombre):
            confirmar = mb.askokcancel('Eliminar cuenta', '¿Estás seguro que quieres eliminar esta cuenta?')
            if confirmar == True:
                lista_cuenta.remove(nombre)
                escribirBD = open(baseDatos, 'wb')
                pickle.dump(lista_cuenta, escribirBD)
                escribirBD.close()
                lista.config(values = lista_cuenta)
                textoDone.set('Cuenta\neliminada')
                done.after(5000, lambda: textoDone.set(''))
                cuenta.set('')
                chequearGenerar()

        def errorNuevo():
            error = mb.showerror('Error', 'Ingresa un nombre válido')
            cuenta.set('')

        def eliminarCuenta():
            nombre = cuenta.get()
            if nombre in lista_cuenta:
                confirmarEliminar(nombre)

        def nuevaCuenta():
            nombre = nombre_cuenta.get()
            if len(nombre) > 2:
                lista_cuenta.append(nombre)
                escribirBD = open(baseDatos, 'wb')
                pickle.dump(lista_cuenta, escribirBD)
                escribirBD.close()
                lista.config(values = lista_cuenta)
                textoDone.set('Cuenta\nagregada')
                done.after(5000, lambda: textoDone.set(''))
                toplevel.destroy()
                lista.set(nombre)
                chequearGenerar()
            else:
                errorNuevo()

        def dialogoCuenta():
            global toplevel 
            toplevel = Toplevel()
            toplevel.title("Crear nueva cuenta")
            mensaje = ttk.Label(toplevel, text = "Escriba nombre de la cuenta")
            mensaje.grid(column = 0, row = 0, sticky = (W, E))
            entrada = ttk.Entry(toplevel, textvariable = nombre_cuenta)
            entrada.grid(column = 0, row = 1, sticky = (W, E))
            boton = ttk.Button(toplevel, text="Guardar", command = nuevaCuenta)
            boton.grid(column = 1, row = 1)


        self.parent.columnconfigure(0, weight = 1)
        self.parent.rowconfigure(0, weight = 1)

        baseDatos = 'cuentas'

        if os.path.isfile('cuentas'):
            leerBD = open(baseDatos, 'rb')
            lista_cuenta = pickle.load(leerBD)
            leerBD.close()
        else:
            lista_cuenta = []

        frame1 = ttk.Frame(self.parent, borderwidth = 1, padding = (4, 4, 4, 4))
        frame1.grid(column = 0, row = 0, sticky = (W, E, S, N))

        frase = ttk.Label(frame1, text = 'Cuenta: ')
        frase.grid(column = 0, row = 1)

        nombre_cuenta = StringVar()
        cuenta = StringVar()
        lista = ttk.Combobox(frame1, textvariable = cuenta)
        lista['values'] = (lista_cuenta)
        lista.state(['readonly'])
        lista.grid(column = 1, row = 1)
        if len(lista_cuenta) > 0:
            lista.current(0)

        agregar_cuenta = ttk.Button(frame1, text = 'Nueva cuenta', command = dialogoCuenta)
        agregar_cuenta.grid(column = 2, row = 1, padx = 2)

        eliminar_cuenta = ttk.Button(frame1, text = 'Eliminar cuenta', command = eliminarCuenta)
        eliminar_cuenta.grid(column = 3, row = 1)

        textoDone = StringVar()
        done = ttk.Label(frame1, textvariable = textoDone)
        done.grid(column = 4, row = 1)

        s1 = ttk.Separator(frame1, orient = HORIZONTAL)
        s1.grid(row = 2, columnspan = 7, sticky = (W, E))

        seleccionPop = ttk.Label(frame1, text = 'Seleccionar criterio para\n clasificar posts')
        seleccionPop.grid(column = 0, row = 3, sticky = (N), pady = 2)

        criterio = StringVar()
        por_impresion = ttk.Radiobutton(frame1, text = 'Impresiones', variable = criterio, value = 'Lifetime_Post_Total_Impressions')
        por_alcance = ttk.Radiobutton(frame1, text = 'Alcance', variable = criterio, value = 'Lifetime_Post_Total_Reach')
        por_engagement = ttk.Radiobutton(frame1, text = 'Tasa de participación', variable = criterio, value = 'engagement_rate')
        por_usuarios = ttk.Radiobutton(frame1, text = 'Usuarios involucrados', variable = criterio, value = 'Lifetime_Engaged_Users')
        #Deja seleccionado el radiobutton por_impresion
        por_impresion.invoke()

        por_impresion.grid(column = 0, row = 4, sticky = W)
        por_alcance.grid(column = 0, row = 5, sticky = W)
        por_engagement.grid(column = 0, row = 6, sticky = W)
        por_usuarios.grid(column = 0, row = 7, sticky = W)

        s2 = ttk.Separator(frame1, orient=VERTICAL)
        s2.grid(column = 1, row = 3, rowspan = 5, sticky = (N, S))

        ingreso_archivos = ttk.Label(frame1, text = 'Seleccione archivos csv')
        ingreso_archivos.grid(column = 2, row = 3, sticky = (E, W), columnspan = 2, pady = 2)

        nivel_post = ttk.Label(frame1, text = 'Nivel post')
        nivel_post.grid(column = 2, row = 4)

        nivel_page = ttk.Label(frame1, text = 'Nivel page')
        nivel_page.grid(column = 2, row = 5)

        filepost = StringVar()
        entrada_post = ttk.Entry(frame1, textvariable = filepost)
        entrada_post.grid(column = 3, row = 4)

        filepage = StringVar()
        entrada_page = ttk.Entry(frame1, textvariable = filepage)
        entrada_page.grid(column = 3, row = 5)

        boton_post = ttk.Button(frame1, text = "..." , command = lambda: seleccionArchivo('post'))
        boton_post.grid(column = 4, row = 4)
        boton_page = ttk.Button(frame1, text = "..." , command = lambda: seleccionArchivo('page'))
        boton_page.grid(column = 4, row = 5)

        s3 = ttk.Separator(frame1, orient = HORIZONTAL)
        s3.grid(row = 8, column = 0, columnspan = 8, sticky = (E, W))

        generar = ttk.Button(frame1, text = 'Generar Informe', state = 'disabled', command = generarInforme)
        generar.grid(row = 9, column = 0, columnspan = 5)

def main():
    root = Tk()
    root.title('Informe de Facebook')
    app = programo(root)
    root.mainloop()

if __name__ == "__main__":
    main()
